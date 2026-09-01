"""Build the public CV download from the reviewed 2026 working draft.

The website is the richer portfolio surface. This document keeps the same
factual source while staying concise enough to be useful as a recruiter-facing
CV. Kahootz detail is intentionally limited to public-safe product leadership
and Product Mole is marked as in progress, not shipped.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

INK = "102A2F"
MUTED = "536467"
TEAL = "0E5E64"
CORAL = "A8432F"
ACID = "D7F276"
PAPER = "EEF0E9"
WHITE = "FFFEF8"
LINE = "C8D0CA"
LIGHT_BLUE = "E7F1F0"


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Calibri", size=11, color=INK, bold=None, italic=None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def set_style_spacing(style, before=0, after=6, line=1.1):
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def get_or_add(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = get_or_add(tc_pr, "w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = get_or_add(tc_pr, "w:tcMar")
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = get_or_add(tc_mar, f"w:{margin}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = get_or_add(tc_pr, "w:tcW")
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA, border_color=LINE):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = get_or_add(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = get_or_add(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = get_or_add(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    borders = get_or_add(tbl_pr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = get_or_add(borders, f"w:{edge}")
        node.set(qn("w:val"), "single" if edge in ("top", "bottom", "insideH") else "nil")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), border_color)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = get_or_add(tr_pr, "w:tblHeader")
    header.set(qn("w:val"), "true")


def set_paragraph_border(paragraph, side="bottom", color=CORAL, size="12", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = get_or_add(p_pr, "w:pBdr")
    border = get_or_add(borders, f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = get_or_add(p_pr, "w:shd")
    shd.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = get_or_add(p_pr, "w:keepNext")
    keep.set(qn("w:val"), "1" if value else "0")


def add_hyperlink(paragraph, text, url, color=TEAL, underline=True, size=9.5):
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(round(size * 2)))
    r_pr.append(size_node)
    if underline:
        underline_node = OxmlElement("w:u")
        underline_node.set(qn("w:val"), "single")
        r_pr.append(underline_node)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, name="Consolas", size=8.5, color=MUTED)


def create_bullet_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = get_or_add(p_pr, "w:numPr")
    ilvl = get_or_add(num_pr, "w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = get_or_add(num_pr, "w:numId")
    num_id_node.set(qn("w:val"), str(num_id))


def add_bullet(document, num_id, text, lead=None):
    paragraph = document.add_paragraph(style="CV Bullet")
    apply_numbering(paragraph, num_id)
    if lead and text.startswith(lead):
        lead_run = paragraph.add_run(lead)
        set_run_font(lead_run, size=9.6, color=INK, bold=True)
        rest = text[len(lead):]
        rest_run = paragraph.add_run(rest)
        set_run_font(rest_run, size=9.6, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=9.6, color=INK)
    return paragraph


def add_section_heading(document, label, title):
    kicker = document.add_paragraph(style="CV Kicker")
    run = kicker.add_run(label.upper())
    set_run_font(run, name="Consolas", size=8, color=CORAL, bold=True)
    set_keep_with_next(kicker)

    heading = document.add_paragraph(style="Heading 1")
    run = heading.add_run(title)
    set_run_font(run, name="Calibri", size=16, color=TEAL, bold=True)
    set_paragraph_border(heading, side="bottom", color=LINE, size="6", space="5")
    set_keep_with_next(heading)
    return heading


def add_role(document, num_id, company, role, dates, context, bullets):
    paragraph = document.add_paragraph(style="CV Role")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    company_run = paragraph.add_run(company)
    set_run_font(company_run, size=12.4, color=TEAL, bold=True)
    separator = paragraph.add_run("  /  ")
    set_run_font(separator, size=11, color=MUTED)
    role_run = paragraph.add_run(role)
    set_run_font(role_run, size=11.5, color=INK, bold=True)
    tab = paragraph.add_run("\t")
    set_run_font(tab, size=10, color=MUTED)
    date_run = paragraph.add_run(dates)
    set_run_font(date_run, name="Consolas", size=8.3, color=MUTED, bold=True)
    set_keep_with_next(paragraph)

    context_para = document.add_paragraph(style="CV Context")
    context_run = context_para.add_run(context)
    set_run_font(context_run, size=9.3, color=MUTED, italic=True)
    set_keep_with_next(context_para)

    for text in bullets:
        add_bullet(document, num_id, text, lead=text.split(":", 1)[0] + ":" if ":" in text else None)


def add_project_table(document):
    rows = [
        (
            "Product Mole",
            "BUILDING / IN PROGRESS\nKahootz AI / product work",
            "A local-first context system for product managers and AI assistants. Building now as part of my Kahootz AI and product work. Not shipped.",
            [("GitHub repo", "https://github.com/simplybenuk/product-mole")],
        ),
        (
            "SourList",
            "LIVE PRODUCT\nSmall SaaS",
            "A focused SaaS to-do list app and a hands-on experiment in taking a small product from idea to a working service.",
            [("sourlist.com", "https://sourlist.com"), ("GitHub profile", "https://github.com/simplybenuk?tab=repositories")],
        ),
        (
            "Wolds Record",
            "PUBLIC REPOS\nContent / product systems",
            "A set of product and content builds, including a lightweight marketing site and supporting social-media workflow tools.",
            [
                ("Marketing repo", "https://github.com/simplybenuk/wolds-record-marketing"),
                ("Social studio repo", "https://github.com/simplybenuk/wolds-record-social-media"),
            ],
        ),
        (
            "The Guide",
            "EXPERIMENT\nAI / interactive story",
            "A mobile-first AI adventure experience.",
            [("GitHub repo", "https://github.com/simplybenuk/the-guide")],
        ),
        (
            "Effective Salary Calculator",
            "LIVE TOOL\nUtility / decision aid",
            "A simple tool for comparing the value of a total salary remuneration package.",
            [
                ("Live tool", "https://salary-renumeration-calculator.vercel.app"),
                ("GitHub repo", "https://github.com/simplybenuk/salary-renumeration-calculator"),
            ],
        ),
        (
            "Playful experiments",
            "INTERACTIVE\nGames / interfaces",
            "The point-and-click puzzle story behind simplyben.co.uk, plus small browser games such as Dodge Balls.",
            [
                ("Site repo", "https://github.com/simplybenuk/simplyben"),
                ("Dodge Balls", "https://dodgeballs.simplyben.co.uk/"),
            ],
        ),
    ]

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    header_row = table.add_row()
    mark_header_row(header_row)
    header_left, header_right = header_row.cells
    set_cell_shading(header_left, INK)
    set_cell_shading(header_right, INK)
    header_left_para = header_left.paragraphs[0]
    header_right_para = header_right.paragraphs[0]
    for paragraph, text in ((header_left_para, "PROJECT / STATUS"), (header_right_para, "WHAT IT IS / LINKS")):
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, name="Consolas", size=7.1, color=WHITE, bold=True)

    for name, meta, description, links in rows:
        row = table.add_row()
        left, right = row.cells
        set_cell_shading(left, LIGHT_BLUE)
        left_para = left.paragraphs[0]
        left_para.paragraph_format.space_after = Pt(3)
        name_run = left_para.add_run(name)
        set_run_font(name_run, size=10.2, color=INK, bold=True)
        meta_para = left.add_paragraph()
        meta_para.paragraph_format.space_after = Pt(0)
        meta_para.paragraph_format.line_spacing = 1.0
        meta_run = meta_para.add_run(meta)
        set_run_font(meta_run, name="Consolas", size=7.1, color=TEAL, bold=True)

        right_para = right.paragraphs[0]
        right_para.paragraph_format.space_after = Pt(3)
        description_run = right_para.add_run(description)
        set_run_font(description_run, size=9.1, color=INK)
        links_para = right.add_paragraph()
        links_para.paragraph_format.space_after = Pt(0)
        for index, (label, url) in enumerate(links):
            if index:
                separator = links_para.add_run("  ·  ")
                set_run_font(separator, name="Consolas", size=7.3, color=MUTED)
            add_hyperlink(links_para, label, url, color=TEAL, size=7.8)
    # Re-apply the exact geometry after rows exist so every cell carries the
    # same DXA width and padding as the table grid.
    set_table_geometry(table, [2450, 6910])
    return table


def set_document_properties(document):
    props = document.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.title = "Ben Whitfield-Heap - CV"
    props.subject = "Product leadership CV and selected work"
    props.comments = ""
    props.keywords = "product leadership, product management, B2B SaaS, govtech, healthtech"


def build_document(output_path: Path):
    document = Document()
    set_document_properties(document)

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    set_style_font(styles["Normal"], size=11, color=INK)
    set_style_spacing(styles["Normal"], before=0, after=6, line=1.1)

    set_style_font(styles["Title"], size=28, color=INK, bold=True)
    set_style_spacing(styles["Title"], before=0, after=4, line=0.95)
    set_style_font(styles["Subtitle"], size=13, color=TEAL)
    set_style_spacing(styles["Subtitle"], before=0, after=10, line=1.0)

    set_style_font(styles["Heading 1"], size=16, color=TEAL, bold=True)
    set_style_spacing(styles["Heading 1"], before=16, after=8, line=1.0)
    set_style_font(styles["Heading 2"], size=13, color=TEAL, bold=True)
    set_style_spacing(styles["Heading 2"], before=12, after=6, line=1.0)
    set_style_font(styles["Heading 3"], size=12, color=INK, bold=True)
    set_style_spacing(styles["Heading 3"], before=8, after=4, line=1.0)

    custom_styles = {
        "CV Kicker": ("Consolas", 8, CORAL, True, False, 0, 3, 1.0),
        "CV Role": ("Calibri", 11, INK, False, False, 8, 2, 1.0),
        "CV Context": ("Calibri", 9.3, MUTED, False, True, 0, 3, 1.0),
        "CV Bullet": ("Calibri", 9.6, INK, False, False, 0, 4, 1.167),
        "CV Small": ("Calibri", 9, MUTED, False, False, 0, 3, 1.0),
    }
    for name, (font, size, color, bold, italic, before, after, line) in custom_styles.items():
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, name=font, size=size, color=color, bold=bold, italic=italic)
        set_style_spacing(style, before=before, after=after, line=line)

    bullet_num_id = create_bullet_numbering(document)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.paragraph_format.space_after = Pt(0)
    header_run = header_para.add_run("BEN WHITFIELD-HEAP  /  PRODUCT LEADER")
    set_run_font(header_run, name="Consolas", size=7.8, color=MUTED, bold=True)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.paragraph_format.space_before = Pt(0)
    footer_run = footer_para.add_run("CV  /  2026  /  ")
    set_run_font(footer_run, name="Consolas", size=8.5, color=MUTED)
    add_page_number(footer_para)

    kicker = document.add_paragraph(style="CV Kicker")
    kicker_run = kicker.add_run("PRODUCT / PRACTICE / PLAY")
    set_run_font(kicker_run, name="Consolas", size=8, color=CORAL, bold=True)

    title = document.add_paragraph(style="Title")
    title_run = title.add_run("Ben Whitfield-Heap")
    set_run_font(title_run, size=28, color=INK, bold=True)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle_run = subtitle.add_run("Product leader / builder")
    set_run_font(subtitle_run, size=13, color=TEAL, bold=True)

    contact = document.add_paragraph(style="CV Small")
    contact.paragraph_format.space_after = Pt(15)
    contact_run = contact.add_run("Lincolnshire, UK  /  Remote-first  /  ")
    set_run_font(contact_run, size=9, color=MUTED)
    add_hyperlink(contact, "benjaminpheap@gmail.com", "mailto:benjaminpheap@gmail.com", size=9)
    separator = contact.add_run("  /  07792 789 316  /  ")
    set_run_font(separator, name="Consolas", size=8.5, color=MUTED)
    add_hyperlink(contact, "linkedin.com/in/simplyben", "https://www.linkedin.com/in/simplyben", size=9)
    set_paragraph_border(contact, side="bottom", color=CORAL, size="12", space="8")

    current = document.add_paragraph(style="CV Small")
    set_paragraph_shading(current, ACID)
    current.paragraph_format.space_before = Pt(3)
    current.paragraph_format.space_after = Pt(15)
    current.paragraph_format.left_indent = Inches(0.12)
    current_run = current.add_run("CURRENTLY  ")
    set_run_font(current_run, name="Consolas", size=8, color=TEAL, bold=True)
    current_role = current.add_run("Head of Product at Kahootz  /  February 2025 to present")
    set_run_font(current_role, size=10, color=INK, bold=True)

    add_section_heading(document, "Profile", "A product leader for the messy middle")
    profile = document.add_paragraph()
    profile.paragraph_format.space_after = Pt(10)
    profile_run = profile.add_run(
        "Product leader with 17 years' experience across B2B SaaS, public-sector, healthtech, and tech-for-good environments. "
        "I shape product strategy, lead discovery, build evidence-led operating models, and help teams align around measurable outcomes. "
        "Currently Head of Product at Kahootz, combining product leadership with hands-on delivery, analytics, enterprise assurance, and practical AI integration."
    )
    set_run_font(profile_run, size=10.5, color=INK)

    add_section_heading(document, "Selected work", "Things I have built, shaped, and kept curious")
    project_intro = document.add_paragraph(style="CV Small")
    project_intro.paragraph_format.space_after = Pt(8)
    project_intro_run = project_intro.add_run("Status is part of the description. Product Mole is building / not shipped.")
    set_run_font(project_intro_run, name="Consolas", size=8.2, color=MUTED)
    add_project_table(document)

    add_section_heading(document, "Experience", "Product leadership across SaaS and public services")
    experience_intro = document.add_paragraph(style="CV Small")
    experience_intro.paragraph_format.space_after = Pt(11)
    experience_intro_run = experience_intro.add_run("Strategy, discovery, delivery, analytics, and the systems that help teams make better decisions.")
    set_run_font(experience_intro_run, size=9.5, color=MUTED)

    add_role(
        document,
        bullet_num_id,
        "Kahootz",
        "Head of Product",
        "February 2025 to present",
        "Secure B2B collaboration SaaS platform",
        [
            "Introduced clearer ownership of roadmap, prioritisation, and product decision-making across the organisation.",
            "Introduced a development portfolio and reporting model across product development, internal work, client-facing work, support, and bug fixing.",
            "Building Product Mole, a local, privacy-conscious AI insight-synthesis and measurement workflow. In progress, not shipped.",
        ],
    )
    add_role(
        document,
        bullet_num_id,
        "Methods",
        "Senior Product Manager",
        "May 2024 to January 2025",
        "Digital agency specialising in public-sector innovation",
        [
            "Led the Department for Transport Respond transformation, creating a three-year roadmap toward automation and identifying AI-powered opportunities for policy drafting, testing, and insight.",
            "Managed a multidisciplinary team of five across product, design, and research, and introduced engagement and satisfaction baselines using Azure Application Insights.",
            "Extended the initial commission from two to nine months by building client confidence through actionable insight and securing additional design resource.",
        ],
    )
    add_role(
        document,
        bullet_num_id,
        "dxw",
        "Senior Product Manager",
        "February 2023 to May 2024",
        "Employee-owned digital agency building technology for good",
        [
            "Digitised the Ministry of Justice Community Accommodation Tier 2 service from 0 to 1, increasing applications by over 100%, reducing digital consent capture from hours to minutes, and integrating with internal MoJ systems.",
            "Established product strategy and reporting baselines across legal, information governance, and third-party suppliers.",
        ],
    )

    document.add_page_break()
    add_section_heading(document, "Experience / continued", "NHS England")
    add_role(
        document,
        bullet_num_id,
        "NHS England",
        "Senior Product Manager",
        "October 2016 to February 2023",
        "FutureNHS collaboration platform and Local Health and Care Records programme",
        [
            "Scaled FutureNHS from 8,000 to more than 50,000 monthly active users in three years by defining product vision and growth strategy for an enterprise collaboration service used across the health and care sector.",
            "Changed the registration process and reduced registration support tickets by 50%; supported a service desk handling 2,500+ tickets per month with a 15-minute average first reply and 90% satisfaction.",
            "Secured £1.5m investment for an NHS-owned open-source collaboration platform and led roadmap, discovery, and external suppliers through service assessment to private beta.",
            "Established an online community of more than 200 digital leaders and created an interactive maturity map used by NHS England's CIO.",
        ],
    )

    add_section_heading(document, "Earlier career", "A foundation in delivery and systems")
    earlier = document.add_paragraph()
    earlier.paragraph_format.space_after = Pt(12)
    earlier_run = earlier.add_run(
        "Senior Project Manager, Attercopia (2015 to 2016); Digital Project Manager, CDS / Bailie Group (2014 to 2015); "
        "Project Manager, Capita Customer Management (2013 to 2014); Systems Project Manager, StepChange Debt Charity (2011 to 2013); "
        "Project Analyst, Wm Morrison Supermarkets (2011); Junior Project Manager, Child Maintenance Commission (2009 to 2011)."
    )
    set_run_font(earlier_run, size=9.5, color=INK)

    add_section_heading(document, "Qualifications and tools", "The useful specifics")
    qualification = document.add_paragraph(style="CV Small")
    qualification.paragraph_format.space_after = Pt(7)
    q_label = qualification.add_run("QUALIFICATIONS  ")
    set_run_font(q_label, name="Consolas", size=8, color=TEAL, bold=True)
    q_text = qualification.add_run("Professional Scrum Product Owner I  ·  AgilePM Practitioner  ·  ITIL Foundation  ·  MSP Practitioner  ·  PRINCE2 Practitioner  ·  BSc (Hons) Business Administration, University of Bath")
    set_run_font(q_text, size=9.2, color=INK)

    tools = document.add_paragraph(style="CV Small")
    tools.paragraph_format.space_after = Pt(0)
    tool_label = tools.add_run("TOOLS AND METHODS  ")
    set_run_font(tool_label, name="Consolas", size=8, color=TEAL, bold=True)
    tool_text = tools.add_run(
        "Product discovery, product strategy, roadmapping, product operations, agile delivery, service design, accessibility, outcome measurement, "
        "Jira, Azure DevOps, Azure Application Insights, Google Analytics, Tag Manager, Looker Studio, Power BI, Tableau, GitHub, GitHub Copilot, "
        "Slack, Zendesk, Google Workspace, and Microsoft 365."
    )
    set_run_font(tool_text, size=9.2, color=INK)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Build the public Ben Whitfield-Heap CV DOCX")
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    build_document(args.output)


if __name__ == "__main__":
    main()
